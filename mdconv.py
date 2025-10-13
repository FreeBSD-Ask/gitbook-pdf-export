#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import shutil
import re
import argparse
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Tuple, Optional
from urllib.parse import urlparse, unquote
from pathlib import Path

from weasyprint import HTML
from weasyprint.text.fonts import FontConfiguration

import mistune
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import html as pygments_html_formatter

from ebooklib import epub
from bs4 import BeautifulSoup, NavigableString

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Pt

    HAS_DOCX = True
except ImportError:
    Document = None
    WD_STYLE_TYPE = None
    qn = None
    Pt = None
    HAS_DOCX = False

# 常量定义
BUILD_DIR = 'build'
SUMMARY_FILE = 'SUMMARY.md'
START_HTML = 'start.html'
FINAL_HTML = 'final.html'
FINAL_PDF = 'final.pdf'
FINAL_EPUB = 'final.epub'
FINAL_DOCX = 'final.docx'
IMAGES_DIR = os.path.join(BUILD_DIR, 'images')
CSS_DIR = 'css'


def prepare_build_dir(build_dir: str) -> None:
    """准备构建目录，清理旧文件"""
    if os.path.exists(build_dir):
        shutil.rmtree(build_dir)
    os.makedirs(build_dir, exist_ok=True)
    os.makedirs(IMAGES_DIR, exist_ok=True)


def read_markdown_file(filename: str) -> str:
    """读取 Markdown 文件内容"""
    with open(filename, 'r', encoding='utf-8') as f:
        return f.read()


def convert_local_paths(html_content: str, docdir: str) -> str:
    """处理本地图片路径和链接路径"""
    def replace_src(match):
        val = match.group(1)
        decoded = unquote(val)
        parsed = urlparse(decoded)
        # 本地资源
        if not parsed.scheme:
            if match.group(0).startswith('src="'):
                abs_path = os.path.abspath(os.path.join(docdir, decoded))
                name, ext = os.path.splitext(os.path.basename(abs_path))
                if not re.match(r'^[A-Za-z0-9]+$', name):
                    md5 = hashlib.md5(name.encode('utf-8')).hexdigest()
                    new_name = md5 + ext
                else:
                    new_name = f"{name}{ext}"
                dst = os.path.join(IMAGES_DIR, new_name)
                try:
                    shutil.copy(abs_path, dst)
                except Exception as e:
                    print(f'Error copying {abs_path}: {e}')
                return f'src="images/{new_name}"'
            elif match.group(0).startswith('href="'):
                if decoded.endswith('.md'):
                    anchor = os.path.splitext(decoded)[0]
                    return f'href="#{anchor}"'
                if '.md#' in decoded:
                    return f'href="#{decoded.split("#")[1]}"'
        return match.group(0)

    html_content = re.sub(r'src="([^"]+)"', replace_src, html_content)
    html_content = re.sub(r'href="([^"]+)"', replace_src, html_content)
    return html_content


class CustomRenderer(mistune.HTMLRenderer):
    """自定义 Markdown 渲染器"""
    def heading(self, text, level, raw=None):
        anchor = re.sub(r'[^\w]+', '-', text.lower()).strip('-')
        level = max(1, min(6, level))
        return f'<h{level} id="{anchor}">{text}</h{level}>'

    def block_code(self, code, info=None):
        if info:
            try:
                lexer = get_lexer_by_name(info.strip(), stripall=True)
            except Exception:
                lexer = get_lexer_by_name('text', stripall=True)
        else:
            lexer = get_lexer_by_name('text', stripall=True)
        fmt = pygments_html_formatter.HtmlFormatter()
        return highlight(code, lexer, fmt)

    def strikethrough(self, text):
        return f'<del>{text}</del>'

    def list_item(self, text):
        m = re.match(r'^\s*\[([ xX])\]\s+', text)
        if m:
            chk = 'checked' if m.group(1).lower()=='x' else ''
            content = text[m.end():]
            return f'<li><input type="checkbox" disabled {chk}> {content}</li>'
        return super().list_item(text)


def markdown_to_html(md_text: str, docdir: str) -> str:
    """将 Markdown 转换为 HTML"""
    renderer = CustomRenderer(escape=False)
    # 启用脚注插件
    md = mistune.create_markdown(renderer=renderer, plugins=['table', 'strikethrough', 'footnotes'])
    html = md(md_text)
    return convert_local_paths(html, docdir)


def combine_markdown_to_html(docdir: str, files: List[str]) -> str:
    """合并多个 Markdown 为单个 HTML 字符串，并支持多线程处理"""

    def process_entry(entry: Tuple[int, str]) -> Tuple[int, str]:
        index, fname = entry
        if fname.startswith('rawchaptertext:'):
            title = fname.split(':', 1)[1]
            aid = re.sub(r'[^\w]+', '-', title.lower()).strip('-')
            html = f'<a id="{aid}"></a>\n<h1>{title}</h1>\n'
            return index, html

        path = os.path.join(docdir, fname)
        if not os.path.exists(path):
            print(f'Warning: {path} 不存在，已跳过')
            return index, ''

        anchor = os.path.splitext(fname)[0]
        md = read_markdown_file(path)
        html_fragment = markdown_to_html(md, os.path.dirname(path)) + '\n'
        return index, f'<a id="{anchor}"></a>\n{html_fragment}'

    combined_parts = [''] * len(files)
    with ThreadPoolExecutor(max_workers=max(4, (os.cpu_count() or 1))) as executor:
        futures = [executor.submit(process_entry, (idx, fname)) for idx, fname in enumerate(files)]
        for future in as_completed(futures):
            idx, html = future.result()
            combined_parts[idx] = html

    combined_html = ''.join(combined_parts)
    print('Markdown 文件已并行转换为 HTML。')
    return combined_html


def merge_with_template(template_path: str, html_content: str, out_html: str) -> None:
    """将生成的 HTML 内容插入模板中，避免出现空白首页"""
    with open(template_path, 'r', encoding='utf-8') as tf:
        soup = BeautifulSoup(tf, 'html.parser')

    body = soup.body
    if body is None:
        raise ValueError('模板中缺少 <body>，无法插入内容')

    body.clear()
    fragment = BeautifulSoup(html_content, 'html.parser')
    for node in fragment.contents:
        if isinstance(node, NavigableString) and not node.strip():
            continue
        body.append(node)

    with open(out_html, 'w', encoding='utf-8') as wf:
        wf.write(str(soup))
    print(f'Merged HTML saved to {out_html}')


def generate_epub_with_ebooklib(html_file: str, start_html: str, output_epub: str):
    """使用 EbookLib 从 HTML 生成 EPUB，包含 start.html 和 CSS"""
    # 读取合并后的内容
    with open(html_file, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    book = epub.EpubBook()
    book.set_identifier('id123456')
    book.set_title('My Book')
    book.set_language('zh')

    # 添加 CSS
    css_items = []
    with open(start_html, 'r', encoding='utf-8') as f:
        start_soup = BeautifulSoup(f, 'html.parser')
    for link in start_soup.find_all('link', rel='stylesheet'):
        href = link.get('href')
        css_path = os.path.join(os.path.dirname(start_html), href)
        if os.path.exists(css_path):
            with open(css_path, 'rb') as cssf:
                css_content = cssf.read()
            css_item = epub.EpubItem(
                uid=os.path.basename(href),
                file_name=os.path.join(CSS_DIR, os.path.basename(href)),
                media_type='text/css',
                content=css_content
            )
            book.add_item(css_item)
            css_items.append(css_item)

    # 添加图片资源
    images_dir = os.path.join(os.path.dirname(html_file), 'images')
    if os.path.exists(images_dir):
        for img_name in os.listdir(images_dir):
            img_path = os.path.join(images_dir, img_name)
            if img_name.lower().endswith('.svg'):
                media_type = 'image/svg+xml'
            elif img_name.lower().endswith(('.png',)):
                media_type = 'image/png'
            elif img_name.lower().endswith(('.jpg', '.jpeg')):
                media_type = 'image/jpeg'
            elif img_name.lower().endswith(('.gif',)):
                media_type = 'image/gif'
            else:
                media_type = 'application/octet-stream'
            with open(img_path, 'rb') as img_file:
                img_content = img_file.read()
            img_item = epub.EpubItem(
                uid=img_name,
                file_name=f'images/{img_name}',
                media_type=media_type,
                content=img_content
            )
            book.add_item(img_item)

    # 按 H1/H2 拆分章节，构建层级目录
    chapters = []
    toc_structure = []
    current_section: Optional[List[epub.EpubHtml]] = None

    body = soup.body or soup
    headers = [hdr for hdr in body.find_all(['h1', 'h2']) if hdr.get_text(strip=True)]

    def build_head_markup() -> str:
        head_parts = ['<meta charset="utf-8"/>']
        for css in css_items:
            head_parts.append(
                f'<link rel="stylesheet" type="text/css" href="{css.file_name}" />'
            )
        return ''.join(head_parts)

    for idx, header in enumerate(headers, start=1):
        chap = epub.EpubHtml(
            title=header.get_text(),
            file_name=f'chap_{idx:04d}.xhtml',
            lang='zh'
        )
        content = [header]
        for sib in header.next_siblings:
            if getattr(sib, 'name', None) in ['h1', 'h2']:
                break
            content.append(sib)
        chap.content = '<html xmlns="http://www.w3.org/1999/xhtml"><head>{}</head><body>{}</body></html>'.format(
            build_head_markup(),
            ''.join(str(tag) for tag in content)
        )
        book.add_item(chap)
        chapters.append(chap)

        level = int(header.name[1])
        if level == 1:
            section = epub.Section(header.get_text())
            toc_structure.append((section, [chap]))
            current_section = toc_structure[-1][1]
        elif level == 2 and current_section is not None:
            current_section.append(chap)
        else:
            toc_structure.append(chap)

    if not chapters:
        # 如果没有明确的标题，使用整个 body 作为单章节
        fallback = epub.EpubHtml(title='内容', file_name='chap_0001.xhtml', lang='zh')
        fallback.content = '<html xmlns="http://www.w3.org/1999/xhtml"><head>{}</head><body>{}</body></html>'.format(
            build_head_markup(),
            ''.join(str(child) for child in body.children)
        )
        book.add_item(fallback)
        chapters.append(fallback)
        toc_structure.append(fallback)

    # 定义目录和 spine
    if toc_structure:
        book.toc = tuple(toc_structure)
    else:
        book.toc = tuple(chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ['nav'] + chapters

    epub.write_epub(output_epub, book, {})
    print(f'EPUB generated via EbookLib: {output_epub}')


def generate_docx_from_html(html_file: str, output_docx: str) -> None:
    """根据最终 HTML 生成 Word 文档"""

    if not HAS_DOCX:
        print('python-docx 未安装，跳过 Word 文档生成。')
        return

    with open(html_file, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    body = soup.body or soup
    document = Document()

    normal_style = document.styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Microsoft YaHei'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal_style.font.bold = False

    try:
        document.styles['CodeBlock']
    except KeyError:
        code_style = document.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        code_style.font.size = Pt(10)
        code_style.font.bold = False

    def append_element(el):
        if isinstance(el, NavigableString):
            text = str(el).strip()
            if text:
                paragraph = document.add_paragraph()
                paragraph.add_run(text)
            return

        if el.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = max(0, min(3, int(el.name[1]) - 1))
            document.add_heading(el.get_text(strip=True), level=level)
        elif el.name == 'p':
            paragraph = document.add_paragraph()
            for child in el.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child))
                elif child.name in ['strong', 'b']:
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['em', 'i']:
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.font.name = 'Consolas'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                elif child.name == 'a':
                    paragraph.add_run(child.get_text())
                else:
                    paragraph.add_run(child.get_text())
        elif el.name in ['ul', 'ol']:
            for li in el.find_all('li', recursive=False):
                paragraph = document.add_paragraph(style='List Number' if el.name == 'ol' else 'List Bullet')
                paragraph.add_run(li.get_text(separator=' ', strip=True))
        elif el.name == 'pre':
            code_text = el.get_text()
            paragraph = document.add_paragraph(style='CodeBlock')
            paragraph.add_run(code_text)
        elif el.name == 'table':
            rows = el.find_all('tr')
            if not rows:
                return
            max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
            table = document.add_table(rows=len(rows), cols=max_cols)
            for r_idx, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for c_idx, cell in enumerate(cells):
                    table.cell(r_idx, c_idx).text = cell.get_text(strip=True)
        else:
            for child in el.children:
                append_element(child)

    for element in body.children:
        append_element(element)

    document.save(output_docx)
    print(f'Word 文档已生成: {output_docx}')


def main(docdir: str):
    prepare_build_dir(BUILD_DIR)

    # 处理 SUMMARY.md
    summ = os.path.join(docdir, SUMMARY_FILE)
    with open(summ, 'r', encoding='utf-8') as f:
        content = f.read()
    summary_md = re.sub(r'## (.+?)$', r'[\1](rawchaptertext:\1)', content, flags=re.MULTILINE)
    summary_build = os.path.join(BUILD_DIR, 'summary.md')
    with open(summary_build, 'w', encoding='utf-8') as f:
        f.write(summary_md)

    # 提取文件列表
    files = re.findall(r'\[.*?\]\((.*?)\)', summary_md)

    # 生成合并 HTML
    combined_html_content = combine_markdown_to_html(docdir, files)
    combined_html_path = os.path.join(BUILD_DIR, 'combined.html')
    with open(combined_html_path, 'w', encoding='utf-8') as wf:
        wf.write(combined_html_content)
    print(f'Combined HTML saved to {combined_html_path}')

    final_html = os.path.join(BUILD_DIR, FINAL_HTML)
    merge_with_template(START_HTML, combined_html_content, final_html)

    # 生成 PDF
    print('Generating PDF (this may take a while)...')
    font_conf = FontConfiguration()
    HTML(final_html).write_pdf(
        os.path.join(BUILD_DIR, FINAL_PDF),
        font_config=font_conf,
        pdf_version="1.7",
        pdf_variant="pdf/a-3u",
        metadata=True,
        base_url=os.path.abspath(BUILD_DIR),
        optimize_size=("fonts",),
    )
    print(f'PDF generated: {os.path.join(BUILD_DIR, FINAL_PDF)}')

    # 生成 EPUB（包含 start.html 和 CSS）
    epub_path = os.path.join(BUILD_DIR, FINAL_EPUB)
    generate_epub_with_ebooklib(final_html, START_HTML, epub_path)

    # 生成 Word 文档
    docx_path = os.path.join(BUILD_DIR, FINAL_DOCX)
    generate_docx_from_html(final_html, docx_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='从 Markdown 生成 PDF 和 EPUB')
    parser.add_argument('docdir', type=str, help='文档根目录')
    args = parser.parse_args()
    main(args.docdir)
